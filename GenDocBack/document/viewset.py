import re
from rest_framework import viewsets, status, filters
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import AllowAny, IsAuthenticated
from django_filters.rest_framework import DjangoFilterBackend
from django.shortcuts import get_object_or_404
from django.core.files.base import ContentFile
from django.db import models
import subprocess
import os
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas # Exemple avec ReportLab
from django.core.files.base import ContentFile
import io
from django.http import FileResponse
import tempfile
from docx import Document
from docx2pdf import convert
import platform
import pythoncom
from .models import (
    CategorieTemplate, TemplateDocument, Formulaire,
    Question, TypeDocument, DocumentGenere, ReponseQuestion
)
from .serializers import (
    CategorieTemplateSerializer, TemplateDocumentListSerializer,
    TemplateDocumentCreateSerializer, TemplateDocumentDetailSerializer,
    FormulaireSerializer, QuestionSerializer, TypeDocumentSerializer,
    DocumentGenereListSerializer, DocumentGenereDetailSerializer,
    DocumentGenereCreateSerializer, ReponseQuestionSerializer
)


class CategorieTemplateViewSet(viewsets.GenericViewSet,
                               viewsets.mixins.ListModelMixin,
                               viewsets.mixins.RetrieveModelMixin,
                               viewsets.mixins.CreateModelMixin):

    queryset = CategorieTemplate.objects.all()
    serializer_class = CategorieTemplateSerializer
    authentication_classes = []
    permission_classes = [AllowAny]
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['nom', 'description']
    ordering_fields = ['nom', 'id']
    ordering = ['nom']

    @action(detail=True, methods=['get'])
    def templates(self, request, pk=None):
        """Récupération de tous les templates d'une catégorie"""
        categorie = self.get_object()
        templates = categorie.templates.filter(status=True)
        serializer = TemplateDocumentListSerializer(templates, many=True)
        return Response(serializer.data)


class TemplateDocumentViewSet(viewsets.GenericViewSet,
                              viewsets.mixins.ListModelMixin,
                              viewsets.mixins.RetrieveModelMixin,
                              viewsets.mixins.CreateModelMixin):

    queryset = TemplateDocument.objects.select_related('categorie').prefetch_related(
        'formulaires__questions'
    )
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    filterset_fields = ['categorie', 'status']
    search_fields = ['nom']
    ordering_fields = ['nom', 'date_add']
    ordering = ['-date_add']

    def get_serializer_class(self):
        """Utilisation des différents serializers selon l'action"""
        if self.action == 'retrieve':
            return TemplateDocumentDetailSerializer
        if self.action == 'create':
            return TemplateDocumentCreateSerializer
        return TemplateDocumentListSerializer

    def perform_create(self, serializer):
        """Après création du template : extraction des variables, création du formulaire et des questions"""
        template = serializer.save()
        self._extract_and_create_questions(template)

    def _extract_and_create_questions(self, template):
        """Ouvre le fichier du template, extrait les variables entre {} et crée le formulaire + questions"""
        variables = self._extract_variables(template.fichier)
        if not variables:
            return

        # Créer le formulaire lié au template
        formulaire = Formulaire.objects.create(
            template=template,
            titre=template.nom
        )

        # Créer une question pour chaque variable
        for variable in variables:
            type_champ = self._infer_type(variable)
            # Générer un nom de variable slug à partir du label
            slug = re.sub(r'[^\w]+', '_', variable.strip().lower()).strip('_')
            Question.objects.create(
                formulaire=formulaire,
                label=variable.strip(),
                variable=slug,
                type_champ=type_champ,
                obligatoire=True
            )

    def _extract_variables(self, fichier):
        """Lit le fichier et extrait les mots entre {}"""
        try:
            ext = fichier.name.split('.')[-1].lower()
            if ext == 'docx':
                return self._extract_from_docx(fichier)
            else:
                fichier.seek(0)
                content = fichier.read().decode('utf-8', errors='ignore')
                return list(dict.fromkeys(re.findall(r'\{([^}]+)\}', content)))
        except Exception as e:
            print(f"Erreur extraction variables: {e}")
            return []

    def _extract_from_docx(self, fichier):
        """Extraction des variables depuis un fichier .docx"""
        try:
            fichier.seek(0)
            doc = Document(fichier)
            full_text = ' '.join([p.text for p in doc.paragraphs])
            # Aussi chercher dans les tableaux
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += ' ' + cell.text
            return list(dict.fromkeys(re.findall(r'\{([^}]+)\}', full_text)))
        except Exception as e:
            print(f"Erreur lecture docx: {e}")
            return []

    def _infer_type(self, variable):
        """Déduit le type de champ à partir du nom de la variable"""
        var = variable.lower()
        if any(k in var for k in ['email', 'e-mail', 'mail', 'courriel']):
            return 'email'
        if any(k in var for k in ['date', 'jour', 'naissance']):
            return 'date'
        if any(k in var for k in ['montant', 'prix', 'total', 'nombre', 'quantite',
                                   'numero', 'numéro', 'num', 'téléphone', 'telephone',
                                   'tel', 'age', 'âge', 'salaire', 'taux', 'pourcentage']):
            return 'number'
        # Par défaut : texte (nom, prenom, adresse, ville, etc.)
        return 'text'

    @action(detail=True, methods=['get'])
    def formulaires(self, request, pk=None):
        """Récupération de tous les formulaires d'un template"""
        template = self.get_object()
        formulaires = template.formulaires.all()
        serializer = FormulaireSerializer(formulaires, many=True)
        return Response(serializer.data)

    @action(detail=False, methods=['get'])
    def actifs(self, request):
        """Récupération des templates actifs"""
        templates = self.queryset.filter(status=True)
        serializer = self.get_serializer(templates, many=True)
        return Response(serializer.data)


class FormulaireViewSet(viewsets.GenericViewSet,
                        viewsets.mixins.ListModelMixin,
                        viewsets.mixins.RetrieveModelMixin,
                        viewsets.mixins.CreateModelMixin):

    queryset = Formulaire.objects.select_related('template').prefetch_related('questions')
    serializer_class = FormulaireSerializer
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    filterset_fields = ['template']
    search_fields = ['titre']
    ordering_fields = ['titre', 'date_add']
    ordering = ['-date_add']


class QuestionViewSet(viewsets.GenericViewSet,
                      viewsets.mixins.ListModelMixin,
                      viewsets.mixins.RetrieveModelMixin,
                      viewsets.mixins.CreateModelMixin):

    queryset = Question.objects.select_related('formulaire')
    serializer_class = QuestionSerializer
    filter_backends = [DjangoFilterBackend, filters.OrderingFilter]
    filterset_fields = ['formulaire', 'type_champ', 'obligatoire']
    ordering_fields = ['id']
    ordering = ['id']


class TypeDocumentViewSet(viewsets.GenericViewSet,
                          viewsets.mixins.ListModelMixin,
                          viewsets.mixins.RetrieveModelMixin,
                          viewsets.mixins.CreateModelMixin):

    queryset = TypeDocument.objects.all()
    serializer_class = TypeDocumentSerializer
    filter_backends = [filters.OrderingFilter]
    filterset_fields = ['status']
    ordering = ['nom']

class DocumentGenereViewSet(viewsets.ModelViewSet):
    queryset = DocumentGenere.objects.select_related('template', 'user').order_by('-date_generation')

    def get_serializer_class(self):
        """Définit quel serializer utiliser selon l'action"""
        if self.action == 'create':
            return DocumentGenereCreateSerializer
        elif self.action == 'retrieve':
            return DocumentGenereDetailSerializer
        return DocumentGenereListSerializer

    def _replace_text_in_docx(self, doc, replacements):
        """
        Remplace les variables dans le document Word.
        AMÉLIORÉ : Gère plusieurs formats de variables et le fractionnement des runs.
        """
        def replace_in_text(text, replacements):
            """Fonction utilitaire pour remplacer dans un texte"""
            for var, val in replacements.items():
                # Support de plusieurs formats : {VAR}, {{VAR}}
                patterns = [
                    f"{{{var}}}",           # {Prénom & Nom du candidat}
                    f"{{{{{var}}}}}",       # {{Prénom & Nom du candidat}}
                ]
                for pattern in patterns:
                    if pattern in text:
                        text = text.replace(pattern, str(val))
            return text

        def process_paragraphs(paragraphs):
            """Traite les paragraphes en gérant le fractionnement des runs"""
            for paragraph in paragraphs:
                # 1. Récupérer le texte complet du paragraphe
                full_text = paragraph.text
                
                # 2. Effectuer les remplacements
                new_text = replace_in_text(full_text, replacements)
                
                # 3. Si des changements ont été effectués
                if new_text != full_text:
                    # Vider tous les runs sauf le premier
                    if paragraph.runs:
                        paragraph.runs[0].text = new_text
                        for i in range(1, len(paragraph.runs)):
                            paragraph.runs[i].text = ""

        # Traiter les paragraphes principaux
        process_paragraphs(doc.paragraphs)
        
        # Traiter les tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_paragraphs(cell.paragraphs)

    def _generer_fichier(self, document_obj, reponses_data):
        """
        Génère le fichier final (DOCX ou PDF) en remplaçant les variables.
        AMÉLIORÉ : Meilleur logging et gestion des erreurs.
        """
        try:
            if not document_obj.template.fichier:
                print("[ERREUR] Aucun fichier template trouve")
                return False

            if not os.path.exists(document_obj.template.fichier.path):
                print("[ERREUR] Fichier template introuvable sur le disque")
                return False

            # Ouvrir le document Word template
            doc = Document(document_obj.template.fichier.path)
            print(f"[OK] Template charge : {document_obj.template.fichier.path}")
            
            # Construction du dictionnaire de remplacement
            replacements = {}
            for r in reponses_data:
                try:
                    question = Question.objects.get(id=r['question'])
                    valeur = r['valeur']
                    
                    # Utiliser le label original (tel qu'il apparaît dans le fichier)
                    replacements[question.label] = valeur
                    # Aussi ajouter le slug au cas où le template utilise ce format
                    replacements[question.variable] = valeur
                    print(f"[MAP] {{{question.label}}} -> {valeur}")
                    
                except Question.DoesNotExist:
                    print(f"[WARN] Question ID {r['question']} non trouvee")
                    continue

            # DEBUG : Afficher toutes les variables trouvées dans le document
            all_text = "\n".join([p.text for p in doc.paragraphs])
            variables_in_doc = re.findall(r'\{([^}]+)\}', all_text)
            print(f"[DEBUG] Variables detectees dans le template : {set(variables_in_doc)}")
            print(f"[DEBUG] Variables a remplacer : {list(replacements.keys())}")

            # Effectuer les remplacements
            self._replace_text_in_docx(doc, replacements)
            print("[OK] Remplacements effectues")

            # Sauvegarder dans un dossier temporaire
            with tempfile.TemporaryDirectory() as tmp_dir:
                temp_docx = os.path.join(tmp_dir, 'out.docx')
                doc.save(temp_docx)
                print(f"[OK] Document temporaire sauvegarde : {temp_docx}")

                ext = 'pdf' if document_obj.format == 'pdf' else 'docx'
                final_name = f"document_{document_obj.id}.{ext}"
                final_path = temp_docx

                # Conversion en PDF si nécessaire
                if document_obj.format == 'pdf':
                    final_path = os.path.join(tmp_dir, 'out.pdf')
                    print("[INFO] Conversion en PDF...")
                    
                    if platform.system() == 'Windows':
                        pythoncom.CoInitialize()
                        try:
                            convert(temp_docx, final_path)
                        finally:
                            pythoncom.CoUninitialize()
                    else:
                        subprocess.run([
                            'libreoffice', '--headless', '--convert-to', 'pdf', 
                            '--outdir', tmp_dir, temp_docx
                        ], check=True)
                    
                    print("[OK] Conversion PDF reussie")

                # Sauvegarder le fichier final
                with open(final_path, 'rb') as f:
                    document_obj.fichier.save(final_name, ContentFile(f.read()), save=False)
                
                print(f"[OK] Fichier final sauvegarde : {final_name}")
            
            document_obj.status = 'done'
            document_obj.save()
            return True
            
        except Exception as e:
            print(f"[ERREUR] Erreur generation : {e}")
            import traceback
            traceback.print_exc()
            document_obj.status = 'error'
            document_obj.save()
            return False

    def create(self, request, *args, **kwargs):
        """Crée un nouveau document généré"""
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        
        # 1. Créer l'objet Document (le serializer gère aussi user et réponses)
        document = serializer.save()
        print(f"[OK] Document cree : ID={document.id}")

        template_file = document.template.fichier
        if not template_file or not os.path.exists(template_file.path):
            document.status = 'error'
            document.save()
            return Response(
                {"error": "Fichier template introuvable. Veuillez recharger le template."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # 2. Récupérer les réponses déjà créées par le serializer
        reponses_data = request.data.get('reponses', [])
        print(f"[INFO] Nombre de reponses : {len(reponses_data)}")

        # 3. Générer le fichier
        if self._generer_fichier(document, reponses_data):
            return Response(
                DocumentGenereDetailSerializer(document).data, 
                status=status.HTTP_201_CREATED
            )
        return Response(
            {"error": "Échec lors de la génération du fichier"}, 
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )

    @action(detail=False, methods=['get'], permission_classes=[IsAuthenticated])
    def history(self, request):
        """Historique des documents générés de l'utilisateur connecté."""
        documents = (
            DocumentGenere.objects
            .filter(user=request.user)
            .select_related('template')
            .order_by('-date_generation')
        )
        serializer = DocumentGenereListSerializer(documents, many=True)
        return Response(serializer.data)

    @action(detail=False, methods=['get'], permission_classes=[IsAuthenticated])
    def recent(self, request):
        """Retourne les 5 documents récents de l'utilisateur connecté."""
        documents = (
            DocumentGenere.objects
            .filter(user=request.user)
            .select_related('template')
            .order_by('-date_generation')[:5]
        )
        serializer = DocumentGenereListSerializer(documents, many=True)
        return Response(serializer.data)

    @action(detail=True, methods=['get'])
    def download(self, request, pk=None):
        """Télécharge un document généré"""
        document = self.get_object()
        if document.fichier and os.path.exists(document.fichier.path):
            return FileResponse(
                open(document.fichier.path, 'rb'), 
                as_attachment=True
            )
        return Response({"detail": "Fichier non trouvé"}, status=404) 
    
    
class ReponseQuestionViewSet(viewsets.GenericViewSet,

                             viewsets.mixins.ListModelMixin,
                             viewsets.mixins.RetrieveModelMixin):

    queryset = ReponseQuestion.objects.select_related('document', 'question')
    serializer_class = ReponseQuestionSerializer
    filter_backends = [DjangoFilterBackend, filters.OrderingFilter]
    filterset_fields = ['document', 'question']
    ordering = ['date_add']